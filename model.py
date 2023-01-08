'''
Nome: Biblioteca para alterar csv
Autor: Sueyvid José
'''
import csv


class Model:
    def __init__(self):
        self.code = 'UTF-8'

    def criar_csv(self, nome_do_arquivo, texto:list):
        '''
        Nome do Arquivo: string com a extensão .csv
        Texto: lista de strings
        '''
        with open(nome_do_arquivo, 'w', newline='', encoding=self.code) as csvfile:
            escritor = csv.writer(csvfile, delimiter=',')
            escritor.writerow(texto)

    def adicionar_csv(self, nome_do_arquivo, texto:list, linha='ultima'):
        '''
        Nome do Arquivo: string com a extensão .csv
        Texto: lista de strings
        '''
        if linha == 'ultima':
            with open(nome_do_arquivo, 'a', newline='', encoding=self.code) as csvfile:
                escritor = csv.writer(csvfile, delimiter=',')
                escritor.writerow(texto)
        else:
            arquivo = self.ler_csv(nome_do_arquivo)
            arquivo.insert(linha, texto)
            self.criar_csv(nome_do_arquivo, arquivo[0])
            for i in arquivo[1:]:
                self.adicionar_csv(nome_do_arquivo, i)

    def ler_csv(self, nome_do_arquivo):
        texto = list()
        with open(nome_do_arquivo, newline='', encoding=self.code) as csvfile:
            leitor = csv.reader(csvfile, delimiter=',')
            for row in leitor:
                texto.append(row)
        return texto

    def retorna_linha_csv(self, nome_do_arquivo, linha):
        return self.ler_csv(nome_do_arquivo)[linha]

    def remover_linha_csv(self, nome_do_arquivo, linha):
        texto = self.ler_csv(nome_do_arquivo)
        texto.pop(linha)
        self.criar_csv(nome_do_arquivo, texto[0])
        for i in texto[1:]:
            self.adicionar_csv(nome_do_arquivo, i)
    
    def editar_linha_csv(self, nome_do_arquivo, musica, linha):
        texto = self.ler_csv(nome_do_arquivo)
        texto.pop(linha)
        texto.insert(linha, musica)
        self.criar_csv(nome_do_arquivo, texto[0])
        for i in texto[1:]:
            self.adicionar_csv(nome_do_arquivo, i)

    def adicionar_cabecalho_csv(self, nome_do_arquivo, colunas):
        texto = self.ler_csv(nome_do_arquivo)
        self.criar_csv(nome_do_arquivo, colunas)
        for linha in texto:
            self.adicionar_csv(nome_do_arquivo, linha)
    
    def alterar_cabecalho_csv(self, nome_do_arquivo, colunas):
        self.remover_linha_csv(0)
        self.adicionar_cabecalho_csv(nome_do_arquivo, colunas)
        
    def copiar_conteudo_csv(self, nome_do_arquivo_copiado, nome_do_arquivo_substituido):
        conteudo = self.ler_csv(nome_do_arquivo_copiado)
        self.criar_csv(nome_do_arquivo_substituido, conteudo[0])
        for linha in conteudo[1:]:
            self.adicionar_csv(nome_do_arquivo_substituido, linha)
    
    def reescrever_conteudo_csv(self, nome_do_arquivo, colunas):
        self.criar_csv(nome_do_arquivo, colunas)

'''
Nome: Biblioteca para alterar docx
Autor: Sueyvid José
'''
from docx import Document
from docx.shared import Pt
from docx import enum, shared


class Docx:
    def __init__(self):
        self.document = Document()

    def font_configs(self, name='Calibri', size=12):
        self.style = self.document.styles['Normal']
        font = self.style.font
        font.name = name
        font.size = Pt(size)

    def paragraph_configs(self):
        self.style.paragraph_format.line_spacing = 0.5
        tab_stops = self.style.paragraph_format.tab_stops
        sec = self.document.sections[0]
        margin_end = shared.Inches(
            sec.page_width.inches - (sec.left_margin.inches + sec.right_margin.inches))
        tab_stops = self.style.paragraph_format.tab_stops
        tab_stops.add_tab_stop(margin_end, enum.text.WD_TAB_ALIGNMENT.RIGHT)

    def escrever_linha(self, texto):
        self.document.add_paragraph(texto).add_run()

    def salvar(self, nome):
        self.document.save(f'{nome}.docx')

if __name__ == '__main__':
    m = Model()
    # m.criar_csv('banco_de_dados.csv', ['Slot 1', 'Slot 2', 'Slot 3'])
    # m.adicionar_csv('banco_de_dados.csv', ['Slot 1', 'Slot 2', 'Slot 4'])
    # m.adicionar_csv('banco_de_dados.csv', ['Slot 1', 'Slot 2', 'Slot 5'])
    # print(m.ler_csv('banco_de_dados.csv'))
    # m.remover_linha_csv('banco_de_dados.csv', 0)
    # m.criar_csv(r'C:\Users\suelt\OneDrive\Documentos\codigos\gerenciador_repertorio\dados\teste.txt', ['teste'])
    # m.subir_linha_csv('temp.csv', 2)
    d = Docx()
    d.font_configs('Arial', 18)
    d.paragraph_configs()
    d.escrever_linha('qualquer coisa\tdb')
    d.escrever_linha('qualquer coisa\tcd')
    d.salvar('teste')